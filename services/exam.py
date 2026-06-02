# services/exam.py - 双语智能解析增强版
import re, json, zhon.hanzi
from datetime import datetime, timezone
from docx import Document
from config import Config
from pytz import UTC
from services.db import get_supabase
import logging
logger = logging.getLogger(__name__)

def _split_bilingual(text: str) -> dict:
    """
    智能分割中英双语文本
    返回: {"raw": 原文, "cn": 中文部分, "en": 英文部分, "mode": "mixed|cn|en"}
    """
    if not text: return {"raw": "", "cn": "", "en": "", "mode": "empty"}
    
    text = text.strip()
    # 🔍 策略 1：按中文句号/问号 + 英文大写字母分割（题干常用）
    m = re.match(r'^([\u4e00-\u9fff\s\p{Punctuation}]+?[.!?。！？;；])\s*([A-Z][a-zA-Z\s\p{Punctuation}/\-]+)$', text, re.UNICODE)
    if m:
        return {"raw": text, "cn": m.group(1).strip(), "en": m.group(2).strip(), "mode": "mixed"}
    
    # 🔍 策略 2：按中英文标点密度分割（启发式）
    cn_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    en_chars = sum(1 for c in text if c.isascii() and c.isalpha())
    
    if cn_chars > en_chars * 2:
        return {"raw": text, "cn": text, "en": "", "mode": "cn"}
    elif en_chars > cn_chars * 2:
        return {"raw": text, "cn": "", "en": text, "mode": "en"}
    else:
        # 混合模式：尝试按空格/标点分割段落
        parts = re.split(r'[\s]{2,}|[\.\!\?\;\;。！？]\s+', text)
        cn_part = ' '.join(p for p in parts if sum(1 for c in p if '\u4e00'<=c<='\u9fff') > len(p)*0.3)
        en_part = ' '.join(p for p in parts if sum(1 for c in p if c.isascii() and c.isalpha()) > len(p)*0.3)
        return {
            "raw": text,
            "cn": cn_part.strip() or text,
            "en": en_part.strip() or text,
            "mode": "mixed"
        }

def _parse_options_bilingual(option_text: str) -> dict:
    """解析单个选项的中英内容（如 "24 小时 24 hours"）"""
    # 🔍 策略：按中文数字/单位 + 英文数字/单位分割
    m = re.match(r'^([\u4e00-\u9fff\d\s\-/]+?)\s+([A-Za-z\d\s\-/]+)$', option_text.strip())
    if m:
        return {"raw": option_text, "cn": m.group(1).strip(), "en": m.group(2).strip()}
    
    # 降级：按空格分割，前段为中文，后段为英文
    parts = option_text.strip().split(None, 1)
    if len(parts) == 2:
        return {"raw": option_text, "cn": parts[0], "en": parts[1]}
    
    return {"raw": option_text, "cn": option_text, "en": "", "mode": "unknown"}

def _post_process(questions: list[dict], exam_id: int = 1) -> list[dict]:
    """后处理：重编号 + 关联考试ID"""
    for i, q in enumerate(questions, 1):
        q["num"] = i
        q["exam_id"] = exam_id
    return questions

def _parse_inline_options_zte(text: str, q: dict):
    """
    解析同行选项
    示例：A、Roscoe C. Howard, Jr.   B、Koh Sow Koon   C、黄智敏   D、颜伟
    """
    import re
    
    # 正则匹配：字母 + 中文顿号/英文逗号 + 内容（到下一个选项或结尾）
    # 支持：A、xxx   B、xxx  或  A,xxx   B,xxx
    pattern = r'([A-E])[、,.]\s*([^A-E、,.]+?)(?=\s+[A-E][、,.]|$)'
    matches = re.findall(pattern, text, re.I)
    
    for key, val in matches:
        q["options"][key.upper()] = val.strip()
    
    logger.debug(f"📝 解析选项：{matches}")

def _clean_option_text(text: str) -> str:
    """清洗选项文本：移除可能残留的选项标识符和多余空白"""
    text = re.sub(r'^[A-E][、,.]\s*', '', text.strip())
    return text.strip()

def _finalize_question(q: dict) -> dict:
    """标准化题目数据，确保字段符合预览模板要求"""
    if not q.get("score"):
        q["score"] = 5

    ans = q.get("answer", "").upper()
    # 推断题型
    if not q.get("type") or q.get("type") == "single":
        if ans in ("T", "F"):
            q["type"] = "judge"
        elif len(ans) > 1 and re.match(r'^[A-E]+$', ans):
            q["type"] = "multi"
        elif q.get("options"):
            q["type"] = "single"
        else:
            q["type"] = "single"

    # 确保 content 字段存在（模板使用 {{ q.content }}）
    q["content"] = q.get("content_cn") or q.get("content_raw", "")
    # 清理题干中可能混入的答案/选项标识
    q["content"] = re.sub(r'\s*答案?[-：:]\s*[A-Ea-e]+', '', q["content"], flags=re.I)
    q["content"] = re.sub(r'\s+', ' ', q["content"]).strip()

    # 清洗选项内容
    clean_options = {}
    for k, v in q.get("options", {}).items():
        clean_options[k] = _clean_option_text(v)
    q["options"] = clean_options

    q["confidence"] = 1.0 if clean_options else 0.5
    q["warnings"] = [] if clean_options else ["⚠️ 未识别到选项"]
    return q

def parse_docx_bilingual(file_path: str, exam_id: int = None) -> list[dict]:
    """
    最终版解析器：全局正则提取选项，完美支持 A~I 及判断题默认选项
    """
    from docx import Document
    import re

    doc = Document(file_path)
    questions = []

    # 提取标题（取第一个非空、非题号的段落）
    title = "未命名考试"
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not re.match(r'^\d+[.,、\s]', text):
            title = text[:100]
            break

    current_type = "single"
    current_score = 5

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info(f"📄 共收集 {len(paragraphs)} 个段落")

    i = 0
    while i < len(paragraphs):
        line = paragraphs[i]

        # 题型声明
        if re.search(r'单选.*?每题.*?分', line, re.I):
            current_type, current_score = "single", _extract_score(line, 5)
            logger.info(f"📋 切换题型: 单选, 分值={current_score}")
            i += 1
            continue
        if re.search(r'多选.*?每题.*?分', line, re.I):
            current_type, current_score = "multi", _extract_score(line, 6)
            logger.info(f"📋 切换题型: 多选, 分值={current_score}")
            i += 1
            continue
        if re.search(r'判断.*?每题.*?分', line, re.I):
            current_type, current_score = "judge", _extract_score(line, 4)
            logger.info(f"📋 切换题型: 判断, 分值={current_score}")
            i += 1
            continue

        # 题号行识别：支持 1,  1.  1、 等
        if re.match(r'^\d+[.,、\s]+', line):
            logger.debug(f"🔍 题号行: {line[:50]}")
            block_lines = [line]
            i += 1
            # 收集后续行，直到遇到下一个题号、题型声明
            while i < len(paragraphs):
                nxt = paragraphs[i]
                if re.match(r'^\d+[.,、\s]+', nxt):
                    break
                if re.search(r'单选|多选|判断', nxt, re.I):
                    break
                block_lines.append(nxt)
                i += 1

            # 解析题目区块
            q = _parse_question_block_v2(block_lines, current_type, current_score, exam_id)
            if q:
                q["num"] = len(questions) + 1
                questions.append(q)
                logger.info(f"✅ 解析题目 #{q['num']}: {q['content'][:30]}... 答案={q['answer']} 选项={list(q['options'].keys())}")
            continue
        else:
            i += 1

    logger.info(f"🎉 总共解析 {len(questions)} 道题目")
    return title, _post_process(questions, exam_id)

def _parse_question_block_v2(block_lines, q_type, q_score, exam_id=None):
    import re

    if not block_lines:
        return None

    first_line = block_lines[0]
    q_match = re.match(r'^\d+[.,、\s]+\s*(.+)', first_line)
    if not q_match:
        return None
    content_cn = q_match.group(1).strip()

    rest_lines = block_lines[1:]

    # 提取答案
    answer = ""
    answer_idx = -1
    for idx, line in enumerate(rest_lines):
        if re.match(r'^答案?[-：:\s]', line, re.I):
            ans_match = re.search(r'答案?[-：:\s]*\s*([A-Ia-i]+|True|False|T|F|√|×|正确|错误|对|错)', line, re.I)
            if ans_match:
                raw = ans_match.group(1).upper()
                if raw in ('TRUE', 'T', '√', '正确', '对'):
                    answer = "T"
                elif raw in ('FALSE', 'F', '×', '错误', '错'):
                    answer = "F"
                elif re.match(r'^[A-I]+$', raw):
                    answer = "".join(sorted(set(raw)))
                else:
                    answer = raw
            answer_idx = idx
            break

    if answer_idx != -1:
        option_lines = rest_lines[:answer_idx]
    else:
        option_lines = rest_lines

    # 同行多选项拆分
    expanded_lines = []
    opt_split_pattern = re.compile(r'([A-I]\s{0,2}[,，、.:：])')
    for line in option_lines:
        line = line.strip()
        if not line:
            continue
        parts = opt_split_pattern.split(line)
        current_fragment = ""
        for part in parts:
            if re.match(r'^[A-I]\s{0,2}[,，、.:：]$', part):
                if current_fragment.strip():
                    expanded_lines.append(current_fragment.strip())
                current_fragment = part
            else:
                current_fragment += part
        if current_fragment.strip():
            expanded_lines.append(current_fragment.strip())

    # 提取选项（保持为简单字符串）
    options = {}
    current_opt_key = None
    opt_start_pattern = re.compile(r'^([A-I])\s{0,2}[,，、.:：]\s*(.*)', re.I)

    for line in expanded_lines:
        line = line.strip()
        if not line:
            continue
        opt_match = opt_start_pattern.match(line)
        if opt_match:
            key = opt_match.group(1).upper()
            val = opt_match.group(2).strip()
            if key in options:
                options[key] += " " + val
            else:
                options[key] = val
            current_opt_key = key
        else:
            if current_opt_key and current_opt_key in options:
                options[current_opt_key] += " " + line

    # 清洗选项内容
    for key in options:
        val = options[key]
        val = re.sub(r'\b[A-I]\s{0,2}[,，、.:：]\s*', '', val)
        val = re.sub(r'\s+', ' ', val).strip()
        options[key] = val

    # 判断题选项处理（动态生成）
    if q_type == "judge" and not options:
        if answer == "T":
            options = {"A": "正确 True", "B": ""}
        elif answer == "F":
            options = {"A": "", "B": "错误 False"}
        else:
            options = {"A": "正确 True", "B": "错误 False"}

    # 清理题干中可能混入的选项标识
    options = {k: v for k, v in options.items() if v and v.strip()}

    return {
        "num": 0,
        "content": content_cn,
        "content_raw": content_cn,
        "content_cn": content_cn,
        "content_en": "",            # 不再使用
        "options": options,          # 简单字符串字典
        "answer": answer,
        "score": q_score,
        "type": q_type,
        "exam_id": exam_id
    }

def _extract_score(line: str, default: int) -> int:
    m = re.search(r'(\d+)\s*分', line)
    return int(m.group(1)) if m else default

def auto_grade(answers: dict, exam_id: int) -> dict:
    """自动评分"""
    db = get_supabase()
    total, details = 0, {}
    q_res = db.table("questions").select("*").eq("exam_id", exam_id).execute()
    logger.info(f"📚 评分：考试 {exam_id}，共 {len(q_res.data or [])} 题")
    
    for q in q_res.data or []:
        uid = str(q["id"])
        u_ans_raw = answers.get(f"q_{uid}", "")
        u_ans = str(u_ans_raw).strip().upper()
        q_type = q.get("type", "single")
        correct_ans = str(q.get("answer", "")).upper()
        score = q.get("score", 0)
        correct = False
        
        if q_type == "single":
            correct = (u_ans == correct_ans)
        elif q_type == "multi":
            # 移除可能存在的空格，确保集合比较准确
            u_set = set(u_ans.replace(" ", ""))
            c_set = set(correct_ans.replace(" ", ""))
            correct = (u_set == c_set)
        elif q_type == "judge":
            norm = u_ans
            if norm in ("A", "T", "√", "正确", "对"):
                norm = "T"
            elif norm in ("B", "F", "×", "错误", "错"):
                norm = "F"
            correct_std = correct_ans.replace("√", "T").replace("×", "F").upper()
            correct = (norm == correct_std)
        
        if correct:
            total += score
        details[uid] = {"correct": correct, "score": score if correct else 0}
        
        logger.debug(f"题目 {uid} 类型={q_type} 考生答案={u_ans} 标准答案={correct_ans} 结果={correct} 得分={score if correct else 0}")
    
    return {"total": total, "details": details}

def save_result(user_id, exam_id, answers, total_score, details, customs=None, submit_method='manual', time_used=None):
    """
    保存考试成绩
    """
    db = get_supabase()
    
    # 验证 submit_method
    if submit_method not in ['manual', 'auto']:
        logger.warning(f"无效的 submit_method: {submit_method}, 使用默认值 'manual'")
        submit_method = 'manual'
    
    # 验证 time_used
    if time_used is not None:
        try:
            time_used = int(time_used)
            if time_used < 0:
                time_used = 0
        except (TypeError, ValueError):
            logger.warning(f"无效的 time_used: {time_used}, 设为 None")
            time_used = None

    print(f"[DEBUG] save_result 被调用")
    print(f"  user_id: {user_id}")
    print(f"  exam_id: {exam_id}")
    print(f"  submit_method: {submit_method}")
    print(f"  传入的 time_used: {time_used}")
    
    # 如果没有传入 time_used，尝试从 user_exam_status 计算
    if time_used is None:
        print(f"[DEBUG] time_used 为 None，尝试从数据库计算")
        try:
            # 获取开始时间和提交时间
            status_res = db.table("user_exam_status") \
                .select("started_at, submitted_at") \
                .eq("user_id", user_id) \
                .eq("exam_id", exam_id) \
                .maybe_single() \
                .execute()
            
            if status_res and status_res.data:
                started_at = status_res.data.get('started_at')
                submitted_at = status_res.data.get('submitted_at')
                print(f"[DEBUG] started_at: {started_at}")
                print(f"[DEBUG] submitted_at: {submitted_at}")
                
                # 如果没有 submitted_at，使用当前时间
                if not submitted_at:
                    submitted_at = datetime.now(timezone.utc).isoformat()
                
                if started_at and submitted_at:
                    # 解析时间
                    start_str = started_at.replace('Z', '+00:00') if started_at.endswith('Z') else started_at
                    submit_str = submitted_at.replace('Z', '+00:00') if submitted_at.endswith('Z') else submitted_at
                    
                    start_dt = datetime.fromisoformat(start_str)
                    submit_dt = datetime.fromisoformat(submit_str)
                    
                    # 确保有时区信息
                    if start_dt.tzinfo is None:
                        start_dt = UTC.localize(start_dt)
                    if submit_dt.tzinfo is None:
                        submit_dt = UTC.localize(submit_dt)
                    
                    time_used = int((submit_dt - start_dt).total_seconds())
                    print(f"[DEBUG] 从数据库计算用时: {time_used} 秒")
                    print(f"[save_result] 计算用时: {time_used} 秒 (start={started_at}, submit={submitted_at})")
        except Exception as e:
            print(f"[save_result] 计算用时失败: {e}")
    else:
        print(f"[DEBUG] 使用传入的 time_used: {time_used}")
    
    # 准备插入数据
    result_data = {
        "user_id": user_id,
        "exam_id": exam_id,
        "answers": json.dumps(answers, ensure_ascii=False) if isinstance(answers, dict) else answers,
        "details": json.dumps(details, ensure_ascii=False) if isinstance(details, dict) else details,
        "total_score": total_score,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "submit_method": submit_method,
        "time_used": time_used,
    }

    print(f"[DEBUG] 准备插入的数据: submit_method={submit_method}, time_used={time_used}")
    
    # 添加自定义字段
    for key, value in customs.items():
        if value:
            result_data[key] = value
    
    # 插入数据库
    result = db.table("exam_results").insert(result_data).execute()
    
    print(f"[save_result] 成绩保存成功: user={user_id}, exam={exam_id}, score={total_score}, method={submit_method}, time_used={time_used}")
    return result
